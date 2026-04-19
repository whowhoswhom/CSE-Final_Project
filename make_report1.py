from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    Image,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parent
FIG = ROOT / "figures"
OUT_PDF = ROOT / "CSE547_FinalProject_Report1_Fuentes.pdf"
OUT_DOCX = ROOT / "CSE547_FinalProject_Report1_Fuentes.docx"


PART1_IR = [
    ["Arch A", "880", "41", "0.7663", "0.7251", "0.7168", "0.6520"],
    ["Arch B", "14,600", "62", "0.8702", "0.8070", "0.8507", "0.7563"],
    ["Arch C", "148,200", "64", "0.9552", "0.8616", "0.9535", "0.8340"],
    ["Arch D", "2,459,848", "9", "0.9766", "0.8880", "0.9772", "0.8792"],
]

PART1_RGB = [
    ["Arch A", "880", "24", "0.6249", "0.5988", "0.5785", "0.5384"],
    ["Arch B", "14,600", "64", "0.8393", "0.7959", "0.8230", "0.7665"],
    ["Arch C", "148,200", "70", "0.9468", "0.8909", "0.9453", "0.8857"],
    ["Arch D", "2,459,848", "39", "0.9999", "0.9201", "0.9999", "0.9185"],
]

PART2_IR_L2 = [
    ["L2 = 1e-5", "0.8984", "0.8984"],
    ["L2 = 1e-4", "0.8862", "0.8903"],
    ["L2 = 1e-3", "0.8724", "0.8491"],
    ["L2 = 1e-2", "0.9064", "0.8973"],
]

PART2_IR_DROPOUT = [
    ["Dropout = 0.10", "0.8984", "0.8838"],
    ["Dropout = 0.25", "0.8924", "0.8743"],
    ["Dropout = 0.40", "0.8977", "0.8809"],
    ["Dropout = 0.60", "0.8746", "0.8420"],
]

PART2_IR_AUG = [
    ["Aug L1", "0.9085", "0.8967"],
    ["Aug L2", "0.9263", "0.9227"],
    ["Aug L3", "0.8908", "0.8837"],
    ["Aug L4", "0.8865", "0.8872"],
]

PART2_IR_COMBINED = [
    ["L2 = 1e-5 + Dropout = 0.1", "0.8786", "0.8601"],
    ["L2 = 1e-5 + Aug L2", "0.9129", "0.9035"],
    ["Dropout = 0.1 + Aug L2", "0.8988", "0.8837"],
    ["L2 + Dropout + Aug", "0.9057", "0.8981"],
]

PART2_RGB_L2 = [
    ["L2 = 1e-5", "0.8883", "0.8835"],
    ["L2 = 1e-4", "0.8965", "0.8894"],
    ["L2 = 1e-3", "0.8495", "0.8335"],
    ["L2 = 1e-2", "0.6777", "0.5766"],
]

PART2_RGB_DROPOUT = [
    ["Dropout = 0.10", "0.9144", "0.9099"],
    ["Dropout = 0.25", "0.8898", "0.8778"],
    ["Dropout = 0.40", "0.8801", "0.8602"],
    ["Dropout = 0.60", "0.8453", "0.8162"],
]

PART2_RGB_AUG = [
    ["Aug L1", "0.9094", "0.9059"],
    ["Aug L2", "0.9043", "0.8997"],
    ["Aug L3", "0.8846", "0.8765"],
    ["Aug L4", "0.8447", "0.8231"],
]

PART2_RGB_COMBINED = [
    ["L2 = 1e-4 + Dropout = 0.1", "0.8994", "0.8947"],
    ["L2 = 1e-4 + Aug L1", "0.8818", "0.8721"],
    ["Dropout = 0.1 + Aug L1", "0.9091", "0.9032"],
    ["L2 + Dropout + Aug", "0.8909", "0.8744"],
]

CLASS_DISTRIBUTION = [
    ["Class", "IR patches", "RGB patches"],
    ["car", "31,513", "54,462"],
    ["person", "12,432", "25,030"],
    ["bike", "3,105", "6,304"],
    ["sign", "2,021", "11,358"],
    ["bus", "1,721", "1,629"],
    ["light", "848", "7,075"],
    ["motor", "730", "1,408"],
    ["truck", "560", "1,031"],
    ["Total", "52,930", "108,297"],
]

ARCHITECTURES = [
    ["Arch A", "2 conv layers with 8 filters, global average pooling, linear output", "880"],
    ["Arch B", "Conv 16, Conv 32, Conv 32, global average pooling, linear output", "14,600"],
    ["Arch C", "2x Conv 32, 2x Conv 64, Conv 128, global average pooling, dense head", "148,200"],
    ["Arch D", "2x Conv 64, 2x Conv 128, 2x Conv 256, Conv 512, batch norm, dense head", "2,459,848"],
]

AUGMENTATION_LEVELS = [
    ["Aug L1", "Random horizontal flip only"],
    ["Aug L2", "Flip, rotation 10 degrees, affine translation 0.05"],
    ["Aug L3", "Flip, rotation 20 degrees, affine translation 0.10, color jitter 0.2"],
    ["Aug L4", "Flip, rotation 30 degrees, affine translation 0.15, color jitter 0.3"],
]


def pdf_styles():
    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="TitleCenter",
            parent=styles["Title"],
            alignment=TA_CENTER,
            fontSize=18,
            leading=22,
            spaceAfter=14,
        )
    )
    styles.add(
        ParagraphStyle(
            name="BodyJustify",
            parent=styles["BodyText"],
            alignment=TA_JUSTIFY,
            fontSize=9.5,
            leading=13,
            spaceAfter=8,
        )
    )
    styles.add(
        ParagraphStyle(
            name="SmallCaption",
            parent=styles["BodyText"],
            fontSize=8.5,
            leading=11,
            spaceAfter=12,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CenterSmall",
            parent=styles["BodyText"],
            alignment=TA_CENTER,
            fontSize=9,
            leading=12,
        )
    )
    return styles


def add_pdf_table(story, title, rows, col_widths=None):
    styles = pdf_styles()
    story.append(Paragraph(f"<b>{title}</b>", styles["BodyJustify"]))
    data = [["Setting", "Val Acc", "Val F1"]] + rows
    if len(rows[0]) == 7:
        data = [["Model", "Params", "Best Ep", "Train Acc", "Val Acc", "Train F1", "Val F1"]] + rows
    if col_widths is None:
        col_widths = [2.4 * inch, 1.0 * inch, 1.0 * inch]
        if len(data[0]) == 7:
            col_widths = [0.85 * inch, 1.05 * inch, 0.65 * inch, 0.85 * inch, 0.85 * inch, 0.85 * inch, 0.85 * inch]
    table = Table(data, colWidths=col_widths, repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#d9eaf7")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.black),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f7f7f7")]),
            ]
        )
    )
    story.append(table)
    story.append(Spacer(1, 0.12 * inch))


def add_pdf_generic_table(story, title, headers, rows, col_widths):
    styles = pdf_styles()
    story.append(Paragraph(f"<b>{title}</b>", styles["BodyJustify"]))
    data = [headers] + rows
    table = Table(data, colWidths=col_widths, repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#d9eaf7")),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f7f7f7")]),
            ]
        )
    )
    story.append(table)
    story.append(Spacer(1, 0.12 * inch))


def add_pdf_figure(story, filename, caption, width=6.5 * inch):
    styles = pdf_styles()
    path = FIG / filename
    if not path.exists():
        story.append(Paragraph(f"<b>Missing figure:</b> {filename}", styles["BodyJustify"]))
        return
    img = Image(str(path))
    scale = width / img.imageWidth
    img.drawWidth = width
    img.drawHeight = img.imageHeight * scale
    max_height = 5.1 * inch
    if img.drawHeight > max_height:
        scale = max_height / img.drawHeight
        img.drawWidth *= scale
        img.drawHeight *= scale
    story.append(img)
    story.append(Paragraph(caption, styles["SmallCaption"]))


def header_footer(canvas, doc):
    canvas.saveState()
    canvas.setFont("Helvetica", 8)
    canvas.drawString(0.6 * inch, 10.55 * inch, "CSE 547 - Deep Learning Algorithms and Methods | Final Project Report 1")
    canvas.drawRightString(7.9 * inch, 10.55 * inch, "Jose Fuentes | University of Louisville")
    canvas.drawRightString(7.9 * inch, 0.45 * inch, f"Page {doc.page}")
    canvas.restoreState()


def build_pdf():
    styles = pdf_styles()
    doc = SimpleDocTemplate(
        str(OUT_PDF),
        pagesize=letter,
        rightMargin=0.6 * inch,
        leftMargin=0.6 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.65 * inch,
    )
    story = []

    story.append(Paragraph("University of Louisville | Speed School of Engineering", styles["CenterSmall"]))
    story.append(Spacer(1, 0.18 * inch))
    story.append(Paragraph("Final Project Report 1", styles["TitleCenter"]))
    story.append(Paragraph("CNN Architectures, Regularization, and Data Augmentation for RGB and Infrared ADAS Object Recognition", styles["TitleCenter"]))
    story.append(Spacer(1, 0.1 * inch))
    meta = [
        ["Course", "CSE 547 - Deep Learning Algorithms and Methods"],
        ["Section", "CSE 547_51 | Spring 2026"],
        ["Instructor", "Dr. Hichem Frigui"],
        ["Student", "Jose Fuentes"],
        ["Due Date", "April 12, 2026"],
    ]
    table = Table(meta, colWidths=[1.3 * inch, 4.6 * inch])
    table.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), 0.25, colors.grey), ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"), ("FONTSIZE", (0, 0), (-1, -1), 9)]))
    story.append(table)
    story.append(Spacer(1, 0.25 * inch))
    story.append(Paragraph("<b>Abstract.</b> This report evaluates RGB and infrared object recognition on the FLIR ADAS v2 dataset for eight autonomous-driving classes: bike, bus, car, person, sign, motor, light, and truck. Part 1 compares four CNN architectures per modality with approximately 10x parameter growth between consecutive models. Part 2 selects a modestly overfitting model per modality and evaluates L2 regularization, dropout, data augmentation, and combined regularization settings. The best raw Part 1 model was Arch D for both modalities, with RGB reaching Val F1 = 0.9185 and IR reaching Val F1 = 0.8792. In Part 2, IR benefited most from moderate data augmentation, reaching Val F1 = 0.9227, while RGB benefited most from light dropout on Arch C, reaching Val F1 = 0.9099.", styles["BodyJustify"]))

    story.append(Paragraph("1. Dataset and Experimental Setup", styles["Heading1"]))
    story.append(Paragraph("The experiments use extracted object patches from the FLIR ADAS v2 dataset. The IR manifest contains 52,930 thermal patches and the RGB manifest contains 108,297 RGB patches. Each patch is labeled as one of eight classes. To prevent leakage from nearby frames and repeated scene conditions, training and validation splits were assigned at the video level for each modality.", styles["BodyJustify"]))
    story.append(Paragraph("All patch classifiers used 64 x 64 RGB inputs, Adam optimization with learning rate 1e-3, cross-entropy loss, batch size 256, and weighted F1 as the primary validation metric. Training was configured for at least 50 epochs with early stopping patience of 10 epochs, but early stopping was only eligible after epoch 50. The maximum epoch count was 70.", styles["BodyJustify"]))
    class_table = Table(CLASS_DISTRIBUTION, colWidths=[1.2 * inch, 1.2 * inch, 1.2 * inch])
    class_table.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#d9eaf7")), ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"), ("GRID", (0, 0), (-1, -1), 0.25, colors.grey), ("FONTSIZE", (0, 0), (-1, -1), 8)]))
    story.append(Paragraph("<b>Table 1. Patch distribution by modality.</b>", styles["BodyJustify"]))
    story.append(class_table)
    story.append(PageBreak())

    story.append(Paragraph("2. Part 1 - CNN Architecture Comparison", styles["Heading1"]))
    story.append(Paragraph("Four CNN architectures were trained per modality. The parameter counts were 880, 14,600, 148,200, and 2,459,848, satisfying the required 10x growth rule between consecutive architectures.", styles["BodyJustify"]))
    add_pdf_generic_table(
        story,
        "Table 2. CNN architecture summary.",
        ["Model", "Architecture summary", "Parameters"],
        ARCHITECTURES,
        [0.75 * inch, 4.6 * inch, 1.05 * inch],
    )
    add_pdf_table(story, "Table 3. Part 1 IR architecture results.", PART1_IR)
    add_pdf_figure(story, "part1_ir_architectures.png", "Figure 1. IR architecture comparison. Validation performance improves as model capacity increases, with Arch D producing the best IR weighted F1.")
    story.append(Paragraph("For IR, model capacity was strongly associated with validation performance. Arch A underfit, while Arch B and Arch C improved both accuracy and weighted F1. Arch D achieved the best IR validation result with Val Acc = 0.8880 and Val F1 = 0.8792, although it also showed a clear train-validation gap.", styles["BodyJustify"]))
    story.append(PageBreak())

    add_pdf_table(story, "Table 4. Part 1 RGB architecture results.", PART1_RGB)
    add_pdf_figure(story, "part1_rgb_architectures.png", "Figure 2. RGB architecture comparison. The largest architecture, Arch D, achieved the strongest RGB validation accuracy and weighted F1.")
    story.append(Paragraph("For RGB, the same trend was observed: increasing model capacity improved validation performance. Arch D achieved the strongest Part 1 RGB result with Val Acc = 0.9201 and Val F1 = 0.9185. RGB therefore outperformed IR in the best raw architecture comparison. However, Arch D showed a larger train-validation gap than Arch C, so Arch C was selected for the RGB Part 2 regularization experiments as the more defensible modest-overfitting base model.", styles["BodyJustify"]))

    story.append(Paragraph("3. Part 2 - IR Regularization and Augmentation", styles["Heading1"]))
    story.append(Paragraph("For IR, Arch D was selected from Part 1 because it produced the best validation performance and had sufficient capacity for meaningful regularization experiments. Four groups of experiments were performed: L2 regularization, dropout, data augmentation, and combined settings.", styles["BodyJustify"]))
    add_pdf_generic_table(
        story,
        "Table 5. Data augmentation level definitions.",
        ["Level", "Transformations"],
        AUGMENTATION_LEVELS,
        [0.8 * inch, 5.6 * inch],
    )
    add_pdf_table(story, "Table 6. IR L2 regularization results.", PART2_IR_L2)
    add_pdf_figure(story, "part2a_ir_l2.png", "Figure 3. IR L2 regularization sweep.")
    add_pdf_table(story, "Table 7. IR dropout results.", PART2_IR_DROPOUT)
    add_pdf_figure(story, "part2b_ir_dropout.png", "Figure 4. IR dropout sweep.")
    story.append(PageBreak())

    add_pdf_table(story, "Table 8. IR data augmentation results.", PART2_IR_AUG)
    add_pdf_figure(story, "part2c_ir_augmentation.png", "Figure 5. IR data augmentation sweep. Augmentation level 2 produced the strongest IR result.")
    add_pdf_table(story, "Table 9. IR combined regularization results.", PART2_IR_COMBINED)
    add_pdf_figure(story, "part2d_ir_combined.png", "Figure 6. IR combined regularization sweep.")
    story.append(Paragraph("The strongest IR Part 2 result was Aug L2 with Val Acc = 0.9263 and Val F1 = 0.9227. This improved over the Part 1 IR Arch D baseline by 0.0435 F1. The combined settings did not exceed the best augmentation-only model, suggesting that regularization effects were not additive for this modality.", styles["BodyJustify"]))
    story.append(PageBreak())

    story.append(Paragraph("4. Part 2 - RGB Regularization and Augmentation", styles["Heading1"]))
    story.append(Paragraph("For RGB, Arch C was selected as the base model for Part 2 because it exhibited moderate overfitting while maintaining strong validation performance. This avoided using Arch D, which achieved the best raw Part 1 score but nearly memorized the training set.", styles["BodyJustify"]))
    add_pdf_table(story, "Table 10. RGB L2 regularization results.", PART2_RGB_L2)
    add_pdf_figure(story, "part2a_rgb_l2.png", "Figure 7. RGB L2 regularization sweep.")
    add_pdf_table(story, "Table 11. RGB dropout results.", PART2_RGB_DROPOUT)
    add_pdf_figure(story, "part2b_rgb_dropout.png", "Figure 8. RGB dropout sweep. Dropout = 0.1 produced the best RGB Part 2 result.")
    story.append(PageBreak())

    add_pdf_table(story, "Table 12. RGB data augmentation results.", PART2_RGB_AUG)
    add_pdf_figure(story, "part2c_rgb_augmentation.png", "Figure 9. RGB data augmentation sweep.")
    add_pdf_table(story, "Table 13. RGB combined regularization results.", PART2_RGB_COMBINED)
    add_pdf_figure(story, "part2d_rgb_combined.png", "Figure 10. RGB combined regularization sweep.")
    story.append(Paragraph("The best RGB Part 2 result was dropout = 0.1 with Val Acc = 0.9144 and Val F1 = 0.9099. This improved over the Part 1 Arch C baseline by 0.0242 F1. Mild augmentation also helped, but stronger geometric and color transformations reduced performance. As with IR, combining the individually best options did not outperform the best single regularization strategy.", styles["BodyJustify"]))

    story.append(Paragraph("5. Cross-Modality Discussion", styles["Heading1"]))
    story.append(Paragraph("The Part 1 results show that RGB achieved higher raw validation performance than IR when both modalities used their best CNN architecture. This likely reflects the richer texture, color, and edge detail available in RGB object patches. However, IR responded more strongly to data augmentation in Part 2: moderate augmentation raised IR Val F1 from 0.8792 to 0.9227, exceeding the best RGB Part 2 result on the selected Arch C base.", styles["BodyJustify"]))
    story.append(Paragraph("The strongest regularization strategy differed by modality. IR benefited most from moderate geometric augmentation, which likely helped the model become less sensitive to object position and thermal-shape variation. RGB benefited most from light dropout, while aggressive augmentation reduced validation performance. This suggests that RGB patches already contain discriminative visual details that can be damaged by strong transformations, whereas IR patches benefit from moderate robustness training.", styles["BodyJustify"]))
    story.append(Paragraph("Another important observation is that combining the best regularization settings did not necessarily produce the best model. In both modalities, the best combined option underperformed the best single option. This indicates interaction effects between regularization methods and supports evaluating each method empirically rather than assuming additive improvements.", styles["BodyJustify"]))

    story.append(Paragraph("6. Conclusions", styles["Heading1"]))
    conclusions = [
        "Increasing CNN capacity improved validation performance in both RGB and IR experiments, and the required 10x parameter growth rule was satisfied.",
        "Arch D was the best Part 1 architecture for both modalities, reaching Val F1 = 0.8792 for IR and Val F1 = 0.9185 for RGB.",
        "For IR Part 2, augmentation level 2 was the best overall regularization strategy, reaching Val F1 = 0.9227.",
        "For RGB Part 2, dropout = 0.1 was the best overall regularization strategy on the selected Arch C base, reaching Val F1 = 0.9099.",
        "The best combined regularization settings did not outperform the best individual settings, showing that regularization effects were not simply additive.",
        "The results suggest that RGB is stronger in raw object-patch classification, while IR benefits substantially from carefully selected augmentation.",
    ]
    for item in conclusions:
        story.append(Paragraph(f"- {item}", styles["BodyJustify"]))

    story.append(Paragraph("6.1 Suggestions for Report 2", styles["Heading2"]))
    for item in [
        "Use class-weighted loss or sampling to address the severe class imbalance, especially for truck, motor, and light.",
        "Use the best RGB regularization result and best IR augmentation result as baselines for later transfer learning, autoencoder, and final-performance experiments.",
        "Add confusion matrices and per-class analysis to identify which classes are most affected by modality and class imbalance.",
        "For final blind-test notebooks, provide a clean prediction pipeline that accepts an unlabeled patch folder and emits class predictions.",
    ]:
        story.append(Paragraph(f"- {item}", styles["BodyJustify"]))

    story.append(Paragraph("References", styles["Heading1"]))
    refs = [
        "Frigui, H. (2026). CSE 547 - Deep Learning Algorithms and Methods. University of Louisville.",
        "Teledyne FLIR. FLIR ADAS Dataset v2.",
        "Paszke, A. et al. (2019). PyTorch: An Imperative Style, High-Performance Deep Learning Library.",
        "TorchVision documentation. Image transformations and pretrained model utilities.",
    ]
    for i, ref in enumerate(refs, 1):
        story.append(Paragraph(f"{i}. {ref}", styles["BodyJustify"]))
    story.append(Paragraph("Note: The code, trained checkpoints, generated manifests, and figures referenced in this report are available in the accompanying project directory and notebooks.", styles["BodyJustify"]))

    doc.build(story, onFirstPage=header_footer, onLaterPages=header_footer)


def add_docx_table(doc, title, rows, headers):
    doc.add_paragraph(title).bold = True
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = str(v)


def add_docx_figure(doc, filename, caption):
    path = FIG / filename
    if path.exists():
        doc.add_picture(str(path), width=Inches(6.4))
        p = doc.add_paragraph(caption)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def build_docx():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    title = doc.add_heading("Final Project Report 1", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph("CNN Architectures, Regularization, and Data Augmentation for RGB and Infrared ADAS Object Recognition")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for line in [
        "University of Louisville | Speed School of Engineering",
        "Course: CSE 547 - Deep Learning Algorithms and Methods",
        "Section: CSE 547_51 | Spring 2026",
        "Instructor: Dr. Hichem Frigui",
        "Student: Jose Fuentes",
        "Due Date: April 12, 2026",
    ]:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("Abstract", level=1)
    doc.add_paragraph(
        "This report evaluates RGB and infrared object recognition on the FLIR ADAS v2 dataset for eight autonomous-driving classes: bike, bus, car, person, sign, motor, light, and truck. Part 1 compares four CNN architectures per modality with approximately 10x parameter growth between consecutive models. Part 2 evaluates L2 regularization, dropout, data augmentation, and combined regularization settings. The best raw Part 1 model was Arch D for both modalities, with RGB reaching Val F1 = 0.9185 and IR reaching Val F1 = 0.8792. In Part 2, IR benefited most from moderate data augmentation, reaching Val F1 = 0.9227, while RGB benefited most from light dropout on Arch C, reaching Val F1 = 0.9099."
    )

    doc.add_heading("1. Dataset and Experimental Setup", level=1)
    doc.add_paragraph(
        "The experiments use extracted object patches from the FLIR ADAS v2 dataset. The IR manifest contains 52,930 thermal patches and the RGB manifest contains 108,297 RGB patches. Training and validation splits were assigned at the video level for each modality to reduce leakage from nearby frames and repeated scene conditions."
    )
    add_docx_table(doc, "Table 1. Patch distribution by modality.", CLASS_DISTRIBUTION[1:], CLASS_DISTRIBUTION[0])

    doc.add_heading("2. Part 1 - CNN Architecture Comparison", level=1)
    doc.add_paragraph("Four CNN architectures were trained per modality. The parameter counts were 880, 14,600, 148,200, and 2,459,848, satisfying the required 10x growth rule.")
    add_docx_table(doc, "Table 2. CNN architecture summary.", ARCHITECTURES, ["Model", "Architecture summary", "Parameters"])
    add_docx_table(doc, "Table 3. Part 1 IR architecture results.", PART1_IR, ["Model", "Params", "Best Ep", "Train Acc", "Val Acc", "Train F1", "Val F1"])
    add_docx_figure(doc, "part1_ir_architectures.png", "Figure 1. IR architecture comparison.")
    add_docx_table(doc, "Table 4. Part 1 RGB architecture results.", PART1_RGB, ["Model", "Params", "Best Ep", "Train Acc", "Val Acc", "Train F1", "Val F1"])
    add_docx_figure(doc, "part1_rgb_architectures.png", "Figure 2. RGB architecture comparison.")
    doc.add_paragraph("Arch D achieved the strongest Part 1 result for both modalities. RGB outperformed IR in raw Part 1 validation performance, reaching Val F1 = 0.9185 compared to IR Val F1 = 0.8792.")

    doc.add_heading("3. Part 2 - IR Regularization and Augmentation", level=1)
    doc.add_paragraph("For IR, Arch D was selected from Part 1 because it produced the best validation performance and had sufficient capacity for regularization experiments.")
    add_docx_table(doc, "Table 5. Data augmentation level definitions.", AUGMENTATION_LEVELS, ["Level", "Transformations"])
    add_docx_table(doc, "Table 6. IR L2 regularization results.", PART2_IR_L2, ["Setting", "Val Acc", "Val F1"])
    add_docx_figure(doc, "part2a_ir_l2.png", "Figure 3. IR L2 regularization sweep.")
    add_docx_table(doc, "Table 7. IR dropout results.", PART2_IR_DROPOUT, ["Setting", "Val Acc", "Val F1"])
    add_docx_figure(doc, "part2b_ir_dropout.png", "Figure 4. IR dropout sweep.")
    add_docx_table(doc, "Table 8. IR augmentation results.", PART2_IR_AUG, ["Setting", "Val Acc", "Val F1"])
    add_docx_figure(doc, "part2c_ir_augmentation.png", "Figure 5. IR data augmentation sweep.")
    add_docx_table(doc, "Table 9. IR combined regularization results.", PART2_IR_COMBINED, ["Setting", "Val Acc", "Val F1"])
    add_docx_figure(doc, "part2d_ir_combined.png", "Figure 6. IR combined regularization sweep.")
    doc.add_paragraph("The strongest IR Part 2 result was Aug L2 with Val Acc = 0.9263 and Val F1 = 0.9227. This improved over the Part 1 IR Arch D baseline by 0.0435 F1.")

    doc.add_heading("4. Part 2 - RGB Regularization and Augmentation", level=1)
    doc.add_paragraph("For RGB, Arch C was selected because it exhibited moderate overfitting while maintaining strong validation performance. Arch D produced the best Part 1 score but nearly memorized the training set.")
    add_docx_table(doc, "Table 10. RGB L2 regularization results.", PART2_RGB_L2, ["Setting", "Val Acc", "Val F1"])
    add_docx_figure(doc, "part2a_rgb_l2.png", "Figure 7. RGB L2 regularization sweep.")
    add_docx_table(doc, "Table 11. RGB dropout results.", PART2_RGB_DROPOUT, ["Setting", "Val Acc", "Val F1"])
    add_docx_figure(doc, "part2b_rgb_dropout.png", "Figure 8. RGB dropout sweep.")
    add_docx_table(doc, "Table 12. RGB augmentation results.", PART2_RGB_AUG, ["Setting", "Val Acc", "Val F1"])
    add_docx_figure(doc, "part2c_rgb_augmentation.png", "Figure 9. RGB data augmentation sweep.")
    add_docx_table(doc, "Table 13. RGB combined regularization results.", PART2_RGB_COMBINED, ["Setting", "Val Acc", "Val F1"])
    add_docx_figure(doc, "part2d_rgb_combined.png", "Figure 10. RGB combined regularization sweep.")
    doc.add_paragraph("The best RGB Part 2 result was dropout = 0.1 with Val Acc = 0.9144 and Val F1 = 0.9099. Mild augmentation also helped, but stronger augmentation reduced validation performance.")

    doc.add_heading("5. Cross-Modality Discussion", level=1)
    doc.add_paragraph("RGB achieved higher raw validation performance than IR in Part 1, likely due to richer texture, color, and edge information. IR benefited more from augmentation in Part 2, suggesting that moderate geometric transformations improved robustness to object placement and thermal-shape variation. In both modalities, combining individually best settings did not outperform the best single strategy.")

    doc.add_heading("6. Conclusions", level=1)
    for item in [
        "Increasing CNN capacity improved validation performance for both modalities.",
        "Arch D was the best Part 1 architecture for both RGB and IR.",
        "IR Part 2 performed best with augmentation level 2, reaching Val F1 = 0.9227.",
        "RGB Part 2 performed best with dropout = 0.1 on Arch C, reaching Val F1 = 0.9099.",
        "The best combined regularization settings did not outperform the best individual settings.",
        "Future work should add class weighting, per-class confusion analysis, transfer learning, autoencoder features, and final blind-test notebooks.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("References", level=1)
    for ref in [
        "Frigui, H. (2026). CSE 547 - Deep Learning Algorithms and Methods. University of Louisville.",
        "Teledyne FLIR. FLIR ADAS Dataset v2.",
        "Paszke, A. et al. (2019). PyTorch: An Imperative Style, High-Performance Deep Learning Library.",
        "TorchVision documentation. Image transformations and pretrained model utilities.",
    ]:
        doc.add_paragraph(ref, style="List Number")
    doc.add_paragraph("Note: The code, trained checkpoints, generated manifests, and figures referenced in this report are available in the accompanying project directory and notebooks.")
    doc.save(str(OUT_DOCX))


if __name__ == "__main__":
    build_pdf()
    build_docx()
    print(OUT_PDF)
    print(OUT_DOCX)
